#!/usr/bin/env python3
"""Compartmentalize Templates.docx: detect ~70 templates, add headings/TOC/page breaks."""

from __future__ import annotations

import argparse
import re
import sys
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


BUILD_RE = re.compile(r"^Build (?:a|an|the)\b", re.I)
CREATE_RE = re.compile(r"^Create (?:a|an|the)\b", re.I)
NUMBERED_RE = re.compile(r"^(\d+)\)\s*(.+?)\s*:?\s*$")

SUBCREATE_SKIP = re.compile(
    r"(css class|@keyframes|component|`<main|\.liquid-glass|reusable|"
    r"ShinyText|hero section container with these exact classes|"
    r"absolutely positioned navbar|content wrapper positioned)",
    re.I,
)

LABEL_NAME_OVERRIDES = {
    "ai workflow agents": "Axon",
    "bold studio": "VANGUARD",
    "interactive discovery": "Lithos",
    "digital director": "Grilled Pixels",
    "mind body healing": "Vibrant Wellness",
    "securify data security": "securify",
    "innovation lab": "VortxLab Creations",
    "fun 404 page": "TinyTrails",
    "wellness devicex": "Measured",
    "bloom ai": "Bloom",
    "stillmind": "Lumora",
    "portal (lock down your passwords)": "Password Manager",
}

# Top-level pre-numbered template labels in the source doc (colon-terminated).
ALLOWED_EARLY_LABELS = frozenset(
    {
        *LABEL_NAME_OVERRIDES.keys(),
        "velorah",
        "aetheris voyage",
        "portal",
        "vex ventures",
        "aethera studio",
        "asme",
        "rivr",
        "power ai",
        "celestrial renewal",
        "luminex",
    }
)

# Boundaries that do not match the colon-label pattern.
EXTRA_BOUNDARIES: list[tuple[int, str, str, str]] = [
    (0, "doc_title", "Templates", "Boomerang"),
    (1217, "label", "Portal (Lock down your passwords)", "Password Manager"),
]

RECREATE_RE = re.compile(r"^Recreate this\b", re.I)
SKIP_PROMPT_PREFIX = re.compile(r"^(?:Build Prompt|PROMPT|Prompt):", re.I)
NUMBERED_TEMPLATE_MIN = 27


@dataclass
class TemplateBlock:
    index: int
    start_para: int
    end_para: int
    name: str
    category: str
    boundary_kind: str
    boundary_label: str


def extract_name_from_prompt(text: str) -> str:
    patterns = [
        r'called ["\u201c]([^"\u201d]+)["\u201d]',
        r"called \*\*([^*]+)\*\*",
        r'for ["\u201c]([^"\u201d]+)["\u201d]',
        r"named ([A-Z][A-Za-z0-9 .'-]+)",
        r'for "([^"]+)"',
        r"for \*\*([^*]+)\*\*",
        r'^["\u201c]([^"\u201d]+)["\u201d]',
    ]
    for pat in patterns:
        m = re.search(pat, text)
        if m:
            return m.group(1).strip().rstrip(".")
    trimmed = text[:70].strip()
    for prefix in ("Build a ", "Build an ", "Build the ", "Create a ", "Create an ", "Create the "):
        if trimmed.lower().startswith(prefix.lower()):
            return trimmed[len(prefix) :].split(".")[0].strip()
    return trimmed


def _normalize_prompt_text(text: str) -> str:
    return text.lstrip("*").strip()


def _find_opening_prompt(doc: Document, idx: int, limit: int = 12) -> tuple[int, str] | None:
    for j in range(idx + 1, min(idx + limit, len(doc.paragraphs))):
        text = doc.paragraphs[j].text.strip()
        if not text or SKIP_PROMPT_PREFIX.match(text):
            continue
        normalized = _normalize_prompt_text(text)
        if BUILD_RE.match(normalized) and len(normalized) >= 40:
            return j, normalized
        if (
            CREATE_RE.match(normalized)
            and len(normalized) >= 40
            and not SUBCREATE_SKIP.search(normalized)
        ):
            return j, normalized
        if RECREATE_RE.match(normalized):
            return j, normalized
    return None


def _is_early_template_label(text: str) -> bool:
    if not text.endswith(":"):
        return False
    if len(text) > 85 or text.startswith("#") or "://" in text:
        return False
    key = text.rstrip(":").strip().lower()
    return key in ALLOWED_EARLY_LABELS


def _name_from_label(label: str, prompt: str = "") -> str:
    key = label.rstrip(":").strip().lower()
    if key in LABEL_NAME_OVERRIDES:
        return LABEL_NAME_OVERRIDES[key]
    return label.rstrip(":").strip()


def detect_boundaries(doc: Document) -> list[tuple[int, str, str, str]]:
    """Return (para_idx, kind, label, name) sorted by index.

    Template boundaries are template-name labels (e.g. ``Velorah:``) or numbered
    section headers (``27) Sentinel Ai:``). Build/Create prompt paragraphs are
    content inside a template block, never boundaries themselves.
    """
    raw: list[tuple[int, str, str, str]] = list(EXTRA_BOUNDARIES)

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        numbered = NUMBERED_RE.match(text)
        if numbered and int(numbered.group(1)) >= NUMBERED_TEMPLATE_MIN:
            name = numbered.group(2).strip().rstrip(":")
            raw.append((i, "numbered", text.rstrip(":"), name))
            continue

        if not _is_early_template_label(text):
            continue

        if not _find_opening_prompt(doc, i):
            continue

        name = _name_from_label(text)
        raw.append((i, "label", text.rstrip(":"), name))

    deduped: list[tuple[int, str, str, str]] = []
    for item in sorted(raw, key=lambda x: x[0]):
        if deduped and item[0] == deduped[-1][0]:
            continue
        deduped.append(item)

    return deduped


def infer_category(name: str, opening_text: str) -> str:
    blob = f"{name} {opening_text}".lower()
    rules = [
        (r"404|error page", "Error Page"),
        (r"portfolio|creative studio", "Portfolio"),
        (r"wellness|health|medical|supplement", "Health & Wellness"),
        (r"saas|software|automation|ai platform|dashboard", "SaaS / Tech"),
        (r"restaurant|food|cafe|culinary", "Food & Hospitality"),
        (r"travel|tourism|wander", "Travel"),
        (r"security|cyber|sentinel|shield|securify", "Security"),
        (r"finance|bank|crypto|defi", "Finance"),
        (r"agency|marketing|email", "Marketing / Agency"),
        (r"hero", "Hero Section"),
        (r"landing", "Landing Page"),
    ]
    for pattern, label in rules:
        if re.search(pattern, blob):
            return label
    return "General"


def analyze_document(doc: Document) -> list[TemplateBlock]:
    boundaries = detect_boundaries(doc)
    if not boundaries:
        raise SystemExit("No template boundaries detected.")

    blocks: list[TemplateBlock] = []
    last_para = len(doc.paragraphs) - 1

    for idx, (start, kind, label, name) in enumerate(boundaries):
        end = boundaries[idx + 1][0] - 1 if idx + 1 < len(boundaries) else last_para
        opening = "\n".join(p.text for p in doc.paragraphs[start : min(start + 8, end + 1)])
        blocks.append(
            TemplateBlock(
                index=idx + 1,
                start_para=start,
                end_para=end,
                name=name,
                category=infer_category(name, opening),
                boundary_kind=kind,
                boundary_label=label,
            )
        )
    return blocks


def _add_hyperlink(paragraph, anchor: str, text: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _bookmark_start(paragraph, bookmark_id: int, name: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    paragraph._p.insert(0, start)


def _bookmark_end(paragraph, bookmark_id: int) -> None:
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)


def _append_paragraph_copy_after(after_el, source_para):
    new_p = deepcopy(source_para._p)
    after_el.addnext(new_p)
    return new_p


def _insert_page_break_after(after_el):
    pb_p = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    pb_p.append(run)
    after_el.addnext(pb_p)
    return pb_p


def organize_document(source_path: Path, output_path: Path) -> list[TemplateBlock]:
    source = Document(source_path)
    blocks = analyze_document(source)
    out = Document()

    title = out.add_paragraph("Templates Library")
    title.style = "Title"

    intro = out.add_paragraph(
        f"{len(blocks)} templates organized with Heading 1 bookmarks, a linked table of "
        "contents, and page breaks between templates."
    )
    intro.style = "Subtitle"

    toc = out.add_paragraph("Table of Contents")
    toc.style = "Heading 1"

    for block in blocks:
        anchor = f"tpl_{block.index:03d}"
        label = f"{block.index}. {block.name} — {block.category}"
        p = out.add_paragraph(style="List Paragraph")
        _add_hyperlink(p, anchor, label)

    out.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    bookmark_id = 1
    source_paras = source.paragraphs

    for block in blocks:
        anchor = f"tpl_{block.index:03d}"
        h1 = out.add_paragraph(style="Heading 1")
        _bookmark_start(h1, bookmark_id, anchor)
        run = h1.add_run(f"{block.index}. {block.name}")
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        _bookmark_end(h1, bookmark_id)
        bookmark_id += 1

        meta = out.add_paragraph(style="Normal")
        meta.add_run("Category: ").bold = True
        meta.add_run(block.category)
        meta.add_run("    Boundary: ").bold = True
        meta.add_run(block.boundary_label[:90])

        spacer = out.add_paragraph(style="Normal")
        last_el = spacer._p

        for para_idx in range(block.start_para, block.end_para + 1):
            last_el = _append_paragraph_copy_after(last_el, source_paras[para_idx])

        if block.index < len(blocks):
            last_el = _insert_page_break_after(last_el)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    out.save(output_path)
    return blocks


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--input",
        type=Path,
        default=Path(r"C:\Users\INTEL\Downloads\Templates.docx"),
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path(r"C:\Users\INTEL\Downloads\Templates-organized.docx"),
    )
    parser.add_argument("--analyze-only", action="store_true")
    args = parser.parse_args(argv)

    if not args.input.exists():
        print(f"Input not found: {args.input}", file=sys.stderr)
        return 1

    source = Document(args.input)
    blocks = analyze_document(source)

    print(f"Input:  {args.input}")
    print(f"Output: {args.output}")
    print(f"Paragraphs in source: {len(source.paragraphs)}")
    print(f"Templates found: {len(blocks)}")
    print(f"Boundary heuristics: doc title, template labels (Name:), numbered ({NUMBERED_TEMPLATE_MIN}+)")
    print()
    for block in blocks:
        section = source.paragraphs[block.start_para : block.end_para + 1]
        builds = sum(1 for p in section if BUILD_RE.match(p.text.strip()))
        creates = sum(
            1
            for p in section
            if CREATE_RE.match(p.text.strip()) and not SUBCREATE_SKIP.search(p.text)
        )
        span = block.end_para - block.start_para + 1
        print(
            f"  {block.index:3d}. {block.name} [{block.category}] "
            f"(para {block.start_para}-{block.end_para}, {span} paras, "
            f"{builds} build, {creates} create)"
        )

    if args.analyze_only:
        return 0

    blocks = organize_document(args.input, args.output)
    print(f"\nSaved organized document: {args.output}")
    print(f"Templates compartmentalized: {len(blocks)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
